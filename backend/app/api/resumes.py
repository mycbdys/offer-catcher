import io
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db import get_db
from app.db.models import Resume
from app.api.deps import require_auth
from app.core.config import settings
from app.schemas.resume import ResumeCreate, ResumeUpdate, ResumeResponse

router = APIRouter()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber (best Chinese support)."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            texts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            result = "\n".join(texts)
            if result.strip():
                return result
    except Exception as e:
        print(f"[PDF] pdfplumber failed: {e}, trying PyPDF2...")

    # Fallback to PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n".join(texts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF文件解析失败: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text.strip())
        return "\n".join(texts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"DOCX文件解析失败: {e}")


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
            if text.strip():
                return text
        except Exception:
            pass
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: .{ext}，支持 PDF/DOCX/DOC/TXT")


@router.post("/upload", response_model=ResumeResponse)
async def upload_resume_file(
    file: UploadFile = File(...),
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    file_bytes = await file.read()
    max_size = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    if len(file_bytes) > max_size:
        raise HTTPException(status_code=400, detail=f"文件不能超过{settings.MAX_UPLOAD_SIZE_MB}MB")

    allowed = {"pdf", "docx", "doc", "txt"}
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in allowed:
        raise HTTPException(status_code=400, detail=f"不支持的格式: .{ext}")

    raw_text = extract_text(file.filename, file_bytes)
    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="文件中未检测到文本内容，请确认文件是否包含文字")

    from app.services.resume_parser import _enhanced_extract
    basic_parsed = _enhanced_extract(raw_text)

    resume = Resume(
        user_id=user_id,
        raw_text=raw_text,
        file_name=file.filename,
        skills=basic_parsed.get("skills", []),
        parsed_data=basic_parsed,
        job_preferences={
            "city": basic_parsed.get("preferences", {}).get("expected_city", ""),
            "positions": basic_parsed.get("preferences", {}).get("expected_positions", []),
            "salary_min": basic_parsed.get("preferences", {}).get("expected_salary_min", 0),
            "salary_max": basic_parsed.get("preferences", {}).get("expected_salary_max", 0),
        },
    )
    db.add(resume)
    await db.commit()
    await db.refresh(resume)
    return ResumeResponse.model_validate(resume)


@router.post("", response_model=ResumeResponse)
async def create_resume_text(
    data: ResumeCreate, user_id: str = Depends(require_auth), db: AsyncSession = Depends(get_db)
):
    resume = Resume(user_id=user_id, raw_text=data.raw_text, file_name=data.file_name)
    db.add(resume)
    await db.commit()
    await db.refresh(resume)
    return ResumeResponse.model_validate(resume)


@router.get("", response_model=list[ResumeResponse])
async def list_resumes(user_id: str = Depends(require_auth), db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Resume).where(Resume.user_id == user_id).order_by(Resume.created_at.desc()))
    return [ResumeResponse.model_validate(r) for r in result.scalars().all()]


@router.get("/{resume_id}", response_model=ResumeResponse)
async def get_resume(resume_id: str, user_id: str = Depends(require_auth), db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Resume).where(Resume.id == resume_id, Resume.user_id == user_id))
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return ResumeResponse.model_validate(resume)


@router.put("/{resume_id}", response_model=ResumeResponse)
async def update_resume(
    resume_id: str, data: ResumeUpdate, user_id: str = Depends(require_auth), db: AsyncSession = Depends(get_db)
):
    result = await db.execute(select(Resume).where(Resume.id == resume_id, Resume.user_id == user_id))
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    if data.parsed_data is not None:
        resume.parsed_data = data.parsed_data
    if data.skills is not None:
        resume.skills = data.skills
    if data.job_preferences is not None:
        resume.job_preferences = data.job_preferences

    await db.commit()
    await db.refresh(resume)
    return ResumeResponse.model_validate(resume)


@router.delete("/{resume_id}")
async def delete_resume(resume_id: str, user_id: str = Depends(require_auth), db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Resume).where(Resume.id == resume_id, Resume.user_id == user_id))
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    await db.delete(resume)
    await db.commit()
    return {"ok": True}
